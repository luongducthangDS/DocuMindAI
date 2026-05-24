"""
Playwright-based crawler for congbao.chinhphu.vn (Cổng Công báo Chính phủ).

Strategy:
  1. Get gazette issue URLs from /cong-bao.htm
  2. For each issue, collect document detail URLs + detect type from slug
  3. Visit each document page → grab DOCX CDN link
  4. Download DOCX with requests (browser cookies, SSL verify=False)
  5. Extract text with python-docx → clean → save JSON

Data is 100% public (Official Gazette). Rate: 1.5s per request.
"""

from __future__ import annotations

import io
import json
import re
import time
import warnings
from pathlib import Path
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from loguru import logger

warnings.filterwarnings("ignore", message="Unverified HTTPS request")

BASE_URL = "https://congbao.chinhphu.vn"
CDN_HOST = "g7.cdnchinhphu.vn"
REQUEST_DELAY = 1.5  # seconds between requests — polite crawling

_TYPE_PREFIXES = [
    ("thong-tu-lien-tich-", "thong_tu_lt"),
    ("bo-luat-", "bo_luat"),
    ("phap-lenh-", "phap_lenh"),
    ("nghi-dinh-", "nghi_dinh"),
    ("nghi-quyet-", "nghi_quyet"),
    ("thong-tu-", "thong_tu"),
    ("quyet-dinh-", "quyet_dinh"),
    ("chi-thi-", "chi_thi"),
    ("luat-", "luat"),
    ("van-ban-hop-nhat-", "hop_nhat"),
]


def _detect_doc_type(url: str) -> str:
    slug = url.rstrip("/").split("/")[-1].split(".")[0]
    for prefix, dtype in _TYPE_PREFIXES:
        if slug.startswith(prefix):
            return dtype
    return "other"


def _safe_filename(title: str) -> str:
    safe = re.sub(r"[^\w\s\-]", "_", title, flags=re.UNICODE)
    return safe[:80].strip()


class CongbaoCrawler:
    def __init__(self, output_dir: Path | None = None):
        from src.config import get_settings

        settings = get_settings()
        self._output = output_dir or settings.data_dir / "raw"
        self._output.mkdir(parents=True, exist_ok=True)
        self._session = requests.Session()
        self._session.headers.update({
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Referer": BASE_URL,
        })

    def crawl(
        self,
        max_issues: int = 20,
        doc_types: list[str] | None = None,
        max_docs: int = 300,
        output_dir: Path | None = None,
    ) -> int:
        """
        Main entry point. Returns count of saved documents.
        doc_types: subset of luat, bo_luat, nghi_dinh, thong_tu, quyet_dinh,
                   nghi_quyet, phap_lenh, chi_thi, thong_tu_lt, hop_nhat, other.
                   None = all types.
        """
        from playwright.sync_api import sync_playwright

        if output_dir:
            self._output = output_dir
            self._output.mkdir(parents=True, exist_ok=True)

        saved = 0
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            try:
                context = browser.new_context(
                    user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
                )
                page = context.new_page()

                # Get gazette issue URLs
                issue_urls = self._get_issue_urls(page, max_issues)
                logger.info("Found {} gazette issues to crawl", len(issue_urls))

                for issue_url in issue_urls:
                    if saved >= max_docs:
                        break

                    doc_links = self._get_issue_doc_links(page, issue_url)
                    logger.info("Issue {} — {} documents", issue_url[-30:], len(doc_links))

                    for doc_info in doc_links:
                        if saved >= max_docs:
                            break

                        dtype = doc_info["doc_type"]
                        if doc_types and dtype not in doc_types:
                            continue

                        url = doc_info["url"]
                        fname = _safe_filename(doc_info["title"]) + ".json"
                        if (self._output / fname).exists():
                            logger.debug("Skip existing: {}", fname)
                            continue

                        docx_url = self._get_docx_url(page, url)
                        if not docx_url:
                            logger.warning("No DOCX link at {}", url)
                            continue

                        cookies = {c["name"]: c["value"] for c in context.cookies()}
                        docx_bytes = self._download_docx(docx_url, cookies)
                        if not docx_bytes:
                            continue

                        text = self._docx_to_text(docx_bytes)
                        if len(text) < 100:
                            logger.debug("Too short, skipping: {}", url)
                            continue

                        self._save_doc(
                            title=doc_info["title"],
                            text=text,
                            url=url,
                            doc_type=dtype,
                            metadata=doc_info.get("metadata", {}),
                        )
                        saved += 1
                        logger.info("[{}/{}] Saved: {:.60s}", saved, max_docs, doc_info["title"])
                        time.sleep(REQUEST_DELAY)

            finally:
                browser.close()

        logger.info("Crawl complete: {} documents saved to {}", saved, self._output)
        return saved

    def _get_issue_urls(self, page, max_issues: int) -> list[str]:
        """Get gazette issue page URLs from the listing."""
        page.goto(f"{BASE_URL}/cong-bao.htm", wait_until="networkidle", timeout=30000)
        time.sleep(1)

        links = page.eval_on_selector_all(
            'a[href*="cong-bao-so"]',
            "els => [...new Set(els.map(e => e.href))].filter(h => !h.includes('#'))",
        )
        # Deduplicate and filter to issue pages only
        seen = set()
        result = []
        for link in links:
            if link not in seen and "cong-bao-so" in link:
                seen.add(link)
                result.append(link)
        return result[:max_issues]

    def _get_issue_doc_links(self, page, issue_url: str) -> list[dict]:
        """Get all document links from a gazette issue page."""
        try:
            page.goto(issue_url, wait_until="networkidle", timeout=20000)
            time.sleep(0.8)
        except Exception as exc:
            logger.warning("Failed to load issue {}: {}", issue_url[-40:], exc)
            return []

        raw = page.eval_on_selector_all(
            'a[href*="/van-ban/"][href$=".htm"]',
            """els => [...new Set(els.map(e => e.href))]
               .filter(h => !h.includes('#') && !h.includes('cong-bao.htm'))
               .map(href => {
                   const el = document.querySelector('a[href=\"' + href.replace(location.origin, '') + '\"]');
                   return {href, text: el ? el.innerText.trim() : ''};
               })""",
        )

        results = []
        seen_urls = set()
        for item in raw:
            url = item["href"]
            if url in seen_urls:
                continue
            seen_urls.add(url)

            # Validate domain (SSRF prevention)
            parsed = urlparse(url)
            if parsed.netloc not in ("congbao.chinhphu.vn",):
                continue

            title = item["text"][:200] or url.split("/")[-1]
            dtype = _detect_doc_type(url)
            results.append({"url": url, "title": title, "doc_type": dtype, "metadata": {}})

        return results

    def _get_docx_url(self, page, doc_url: str) -> str | None:
        """Navigate to document page and extract DOCX CDN download URL."""
        # Validate target URL before fetching
        parsed = urlparse(doc_url)
        if parsed.netloc != "congbao.chinhphu.vn":
            logger.error("Refusing non-congbao URL: {}", doc_url)
            return None

        try:
            page.goto(doc_url, wait_until="networkidle", timeout=20000)
            time.sleep(0.5)
        except Exception as exc:
            logger.warning("Failed to load doc page {}: {}", doc_url[-50:], exc)
            return None

        # Find DOCX download link from CDN
        docx_links = page.eval_on_selector_all(
            f'a[href*="{CDN_HOST}"][href*=".docx"]',
            "els => els.map(e => e.href)",
        )
        if docx_links:
            return docx_links[0]

        # Fallback: look for any docx link
        fallback = page.eval_on_selector_all(
            'a[href*=".docx"]',
            "els => els.map(e => e.href)",
        )
        return fallback[0] if fallback else None

    def _download_docx(self, docx_url: str, cookies: dict) -> bytes | None:
        """Download DOCX file from CDN. Returns bytes or None on failure."""
        try:
            resp = self._session.get(
                docx_url,
                cookies=cookies,
                timeout=30,
                verify=False,  # CDN uses self-signed cert on this network
                stream=True,
            )
            resp.raise_for_status()
            content = resp.content
            if len(content) < 500:
                logger.warning("DOCX too small ({} bytes), likely error page", len(content))
                return None
            logger.debug("Downloaded {} bytes", len(content))
            return content
        except Exception as exc:
            logger.warning("DOCX download failed: {}", str(exc)[:80])
            return None

    def fetch_download_page(self, url: str, metadata: dict | None = None) -> dict | None:
        """
        Fetch a specific Congbao download/detail page and extract full text from
        all attached PDFs. This is more reliable for old laws whose public pages
        are split across multiple gazette PDF attachments.
        """
        parsed = urlparse(url)
        if parsed.netloc != "congbao.chinhphu.vn":
            logger.error("Refusing non-congbao URL: {}", url)
            return None

        try:
            resp = self._session.get(url, timeout=30, verify=False)
            resp.raise_for_status()
        except Exception as exc:
            logger.warning("Failed to fetch Congbao download page {}: {}", url, exc)
            return None

        soup = BeautifulSoup(resp.text, "lxml")
        page_title = soup.find("h1")
        title = (
            (metadata or {}).get("title")
            or (page_title.get_text(" ", strip=True) if page_title else "")
            or (soup.title.get_text(" ", strip=True) if soup.title else "congbao_document")
        )

        pdf_links = []
        for a in soup.select("a[href]"):
            href = a.get("href", "")
            text = a.get_text(" ", strip=True).lower()
            if ".pdf" in href.lower() or ".pdf" in text:
                pdf_links.append(href)

        # Preserve order while deduplicating.
        seen = set()
        pdf_links = [x for x in pdf_links if not (x in seen or seen.add(x))]

        if not pdf_links:
            logger.warning("No PDF attachments found at {}", url)
            return None

        parts = []
        for pdf_url in pdf_links:
            text = self._download_pdf_text(pdf_url)
            if text:
                parts.append(text)
            time.sleep(REQUEST_DELAY)

        full_text = "\n\n".join(parts)
        if len(full_text) < 1000:
            logger.warning("Extracted text too short from {} ({} chars)", url, len(full_text))
            return None

        from src.ingestion.cleaner import clean_legal_text

        meta = metadata or {}
        return {
            "title": title,
            "content": clean_legal_text(full_text, doc_title=title),
            "url": url,
            "source": "congbao.chinhphu.vn",
            "doc_type": meta.get("doc_type", _detect_doc_type(url)),
            "so_hieu": meta.get("so_hieu", ""),
            "ngay_ban_hanh": meta.get("ngay_ban_hanh", ""),
        }

    def _download_pdf_text(self, pdf_url: str) -> str:
        """Download an attached PDF and extract all page text."""
        try:
            resp = self._session.get(
                pdf_url,
                timeout=60,
                verify=False,
                headers={"User-Agent": "Mozilla/5.0"},
            )
            resp.raise_for_status()
            content = resp.content
            if not content.startswith(b"%PDF"):
                logger.warning("Attachment is not a PDF: {}", pdf_url[:100])
                return ""

            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [page.extract_text(x_tolerance=1, y_tolerance=3) or "" for page in pdf.pages]
            text = "\n".join(pages)
            logger.info("Extracted {} chars from PDF attachment", len(text))
            return text
        except Exception as exc:
            logger.warning("PDF extraction failed: {}", str(exc)[:120])
            return ""

    def _docx_to_text(self, docx_bytes: bytes) -> str:
        """Extract plain text from DOCX bytes using python-docx."""
        try:
            import docx as docxlib

            doc = docxlib.Document(io.BytesIO(docx_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt = cell.text.strip()
                        if txt and txt not in paragraphs:
                            paragraphs.append(txt)
            return "\n\n".join(paragraphs)
        except Exception as exc:
            logger.warning("DOCX extraction failed: {}", exc)
            return ""

    def _save_doc(
        self,
        title: str,
        text: str,
        url: str,
        doc_type: str,
        metadata: dict,
    ) -> None:
        """Save document as JSON. Idempotent — skips if file exists."""
        from src.ingestion.cleaner import clean_legal_text

        cleaned = clean_legal_text(text, doc_title=title)
        doc = {
            "title": title,
            "content": cleaned,
            "url": url,
            "source": "congbao.chinhphu.vn",
            "doc_type": doc_type,
            "so_hieu": metadata.get("so_hieu", ""),
            "ngay_ban_hanh": metadata.get("ngay_ban_hanh", ""),
        }
        fname = _safe_filename(title) + ".json"
        out_path = self._output / fname
        if out_path.exists():
            return
        out_path.write_text(json.dumps(doc, ensure_ascii=False, indent=2), encoding="utf-8")
        logger.debug("Saved: {}", fname)
